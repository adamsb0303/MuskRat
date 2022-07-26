from email import header
import docx
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
import hashlib

path = r'C:\Users\624409\Desktop\Image Classification\BAH Logo.jpg'
y = hashlib.md5(path.encode('utf-8')).hexdigest()
print(y)


document = Document()
document = docx.Document()
run = document.add_paragraph().add_run()
# The style allows for paragraph adjustments to the font and size
header = document.sections[0].header
header.paragraphs[0].text = "M.U.S.K.R.A.T."

style = document.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = docx.shared.Pt(12)

# document.add_picture(r'C:\Users\624409\Desktop\Image Classification\logo_png.png', width = docx.shared.Inches(1), height = docx.shared.Inches(.5))
# last_paragraph = document.paragraphs[-1]
# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
# MUSKTRAT'S Logo

document.add_heading('Results of Image/Video', 0)

paragraph = document.add_paragraph('Image Submitted: \n')

document.add_picture(r'C:\Users\624409\Desktop\Image Classification\Elon.jpg', width = docx.shared.Inches(1.5), height = docx.shared.Inches(1.5))

paragraph = document.add_paragraph('''MuskRat's model prediction:\n''')

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text('Model 1 = 72%')
# Use all four lines for any one bullet

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text('Model 2 = 84%')

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text('Model 3 = 76%')

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text('Model 4 = 87%')

p = document.add_paragraph()
p.style = 'List Bullet'
r = p.add_run()
r.add_text('Verdict: Real = 79.95%')


document.save('DeepFake Output.docx')